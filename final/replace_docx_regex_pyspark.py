from docx import Document
from io import BytesIO
from pyspark.sql import SparkSession
import os
import re

def find_and_replace_sensitive_info(binary_content, email_formats, credit_card_formats):
    """
    Find and replace sensitive information (email addresses and credit card numbers) 
    in the binary content of a DOCX file.

    Args:
        binary_content (bytes): The binary content of the DOCX file.
        email_formats (list): List of regular expressions for identifying email addresses.
        credit_card_formats (list): List of regular expressions for identifying credit card numbers.

    Returns:
        bytes: The modified binary content.
    """
    docx_bytes = BytesIO(binary_content)
    document = Document(docx_bytes)
    count_changes = 0  # Counter for changes made

    for paragraph in document.paragraphs:
        # Replace email addresses
        for email_regex in email_formats:
            if re.search(email_regex, paragraph.text):
                paragraph.text = re.sub(email_regex, "changed-email", paragraph.text)
                count_changes += 1
        
        # Replace credit card numbers
        for card_regex in credit_card_formats:
            if re.search(card_regex, paragraph.text):
                paragraph.text = re.sub(card_regex, "changed-card", paragraph.text)
                count_changes += 1

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Replace email addresses
                    for email_regex in email_formats:
                        if re.search(email_regex, paragraph.text):
                            paragraph.text = re.sub(email_regex, "changed-email", paragraph.text)
                            count_changes += 1
                    
                    # Replace credit card numbers
                    for card_regex in credit_card_formats:
                        if re.search(card_regex, paragraph.text):
                            paragraph.text = re.sub(card_regex, "changed-card", paragraph.text)
                            count_changes += 1

    modified_docx_bytes = BytesIO()
    document.save(modified_docx_bytes)
    return modified_docx_bytes.getvalue(), count_changes


def save_to_docx(row):
    """
    Save the modified binary content to a DOCX file.

    Args:
        row: Row containing file path, modified binary content, and changes count.
    """
    file_path = row.File
    modified_content = row.ModifiedBinaryContent

    # Specify the output directory for saving modified DOCX files
    output_directory = "/content/o2"

    # Create the output directory if it doesn't exist
    os.makedirs(output_directory, exist_ok=True)

    # Extract the file name from the path
    file_name = os.path.basename(file_path)

    # Construct the output file path
    output_file_path = os.path.join(output_directory, file_name)

    # Save the modified binary content to the output file
    with open(output_file_path, 'wb') as file:
        file.write(modified_content)

# Initialize Spark session
spark = SparkSession.builder.appName("FindAndReplaceInDOCX").getOrCreate()

# Specify the path to the directory containing DOCX files
input_directory = "/content/output"

# Read binary DOCX files directly into a DataFrame
docx_files_df = (
    spark.read.format("binaryFile")
    .option("pathGlobFilter", "*.docx")
    .load(input_directory)
)

# Define email and credit card number formats
email_id_formats = [
    r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}",
]

credit_card_number_formats = [
    r"\b^4[0-9]{12}(?:[0-9]{3})?$\b",
    r"\b^5[1-5][0-9]{14}$\b",
    r"\b^3[47][0-9]{13}$\b",
    r"\b^(6011|65|64[4-9])[0-9]{13}$\b",
    r"\b^3(?:0[0-5]|[68][0-9])[0-9]{11}$\b",
    r"\b^(?:2131|1800|35\d{3})\d{11}$\b",
    r"\b^(62|88)\d{14,17}$\b",
    r"\b(5[1-5][0-9]{2}[0-9]{4}[0-9]{4}[0-9]{4})|4[0-9]{3}[0-9]{4}[0-9]{4}[0-9]{4}|4[0-9]{3}[0-9]{4}[0-9]{4}[0-9]{1}|3[4,7][0-9]{2}[0-9]{6}[0-9]{5}|((6011)[0-9]{4}[0-9]{4}[0-9]{4})|((65)[0-9]{2}[0-9]{4}[0-9]{4}[0-9]{4})|(3[0478][0-9]{2}[0-9]{4}[0-9]{4}[0-9]{2,4})\b"
]

# Function to apply find-and-replace on each row of the DataFrame
def apply_find_and_replace_sensitive_info(row):
    file_path = row.path
    binary_content = row.content
    modified_content, count_changes = find_and_replace_sensitive_info(binary_content, email_id_formats, credit_card_number_formats)
    return file_path, modified_content, count_changes

# Apply find-and-replace on each row of the DataFrame
modified_docx_df = docx_files_df.rdd.map(apply_find_and_replace_sensitive_info).toDF(["File", "ModifiedBinaryContent", "ChangesCount"])

# Save the modified binary content to DOCX files
modified_docx_df.foreach(save_to_docx)

# Show the results (for demonstration purposes, use show wisely)
modified_docx_df.show(truncate=True)

# Stop Spark session
spark.stop()
