from docx import Document
from io import BytesIO
def find_and_replace(document: Document, target_text: str, replacement_text: str):
    """
    Find and replace the specified text in the document.

    Args:
        document (Document): The document object.
        target_text (str): The text to be replaced.
        replacement_text (str): The text to replace the target text with.
    """
    for paragraph in document.paragraphs:
        index = paragraph.text.find(target_text)
        if index != -1:
            for run in paragraph.runs:
                run_index = run.text.find(target_text)
                if run_index != -1:
                    run.text = run.text.replace(target_text, replacement_text)
                    break

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    index = paragraph.text.find(target_text)
                    if index != -1:
                        for run in paragraph.runs:
                            run_index = run.text.find(target_text)
                            if run_index != -1:
                                run.text = run.text.replace(target_text, replacement_text)
                                break


input_docx_file_path = '/home/axat/Downloads/Findal copy copy.docx'

with open(input_docx_file_path, 'rb') as file:
    docx_bytes = BytesIO(file.read())
    document = Document(docx_bytes)

target_text = 'xzat'
replacement_text = 'sumit'

find_and_replace(document, target_text, replacement_text)

with open(input_docx_file_path, 'wb') as file:
    document.save(file)