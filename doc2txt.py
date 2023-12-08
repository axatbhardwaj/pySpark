from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def read_docx_objects(docx_file_path):
    doc = Document(docx_file_path)
    content = ""
    processed_headers = set()
    for section in doc.sections:
        for header in section.header.paragraphs:
            if header.text not in processed_headers:
                content += f'Header: {header.text}\n'
                processed_headers.add(header.text)

    content = ""
    for element in list(doc.iter_inner_content()):
        if isinstance(element, Paragraph):
            content += f'{element.text}\n'
        elif isinstance(element, Table):
            for row in element.rows:
                row_text = ' '.join([cell.text.replace('\n', ' ') for cell in row.cells])
                content += f'{row_text}\n'

    processed_footers = set()
    for section in doc.sections:
        for footer in section.footer.paragraphs:
            if footer.text not in processed_footers:
                content += f'Footer: {footer.text}\n'
                processed_footers.add(footer.text)

    return content


input_docx_file_path = ''

document_content = read_docx_objects(input_docx_file_path)

print(document_content)
